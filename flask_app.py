from flask import Flask, request, render_template, Markup
import os
import xml.etree.ElementTree as ET
import html
import tempfile
import re
from lxml import etree
from urllib.parse import quote

from docx import Document

import pypandoc


app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

@app.route('/')
def index():
    return render_template('index.html')


@app.route('/upload', methods=['POST'])
def upload():
    if 'file' not in request.files:
        return 'No file part'

    file = request.files['file']

    if file.filename == '':
        return 'No selected file'

    filename = os.path.join(UPLOAD_FOLDER, file.filename)
    file.save(filename)

    print("File saved successfully. Starting extraction...")

    equations = extract_math_from_docx(filename)

    return render_template('display.html', equations=equations, filename=file.filename)


def extract_math_from_docx(docx_filename):
    doc = Document(docx_filename)
    equations = []
    current_header = None

    # xml namespace for math
    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
             'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}

    for element in doc.element.body:
        # Parse current XML element
        tree = ET.ElementTree(element)
        print("# Find all oMath elements (Word math objects)")
        
        extracted_header = _extract_header_from_element(tree, nsmap)
        if extracted_header:
            current_header = extracted_header
            print("Current Header:", current_header)

        # Extract math equations
        equations += _extract_equations_from_element(tree, nsmap, current_header)

    return equations


def _extract_header_from_element(tree, nsmap):
    style = tree.find(".//w:pStyle", nsmap)
    if style is not None:
        style_val = style.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
        
        if style_val in ['1', '2', '3', '4', '5']:
            headingtag = '#' * int(style_val) + ' '
            header_text = ''.join(run.text for run in tree.getroot().findall(".//w:t", nsmap))
            return headingtag + header_text
    return None


def _extract_equations_from_element(tree, nsmap, current_header):
    equations = []
    for omath in tree.findall('.//m:oMath', nsmap):
        omath_str = ET.tostring(omath, encoding='unicode')
        latex_code = omath_to_latex_via_docx(omath_str)
        mode = determine_mode(omath, nsmap)
        equations.append((current_header, latex_code, mode))
    return equations


def omath_to_latex_via_docx(omath_str):
    # 문자열로부터 lxml element를 생성
    omath_element = etree.fromstring(omath_str)

    # 임시 Document 객체 생성
    doc = Document()
    
    # oMath 추가
    p = doc.add_paragraph()
    p._element.append(omath_element)
    
    # 임시 docx 파일 저장
    with tempfile.NamedTemporaryFile(delete=True, suffix='.docx') as temp:
        doc.save(temp.name)
        # Pandoc을 사용하여 LaTeX로 변환
        latex_code = pypandoc.convert_file(temp.name, 'latex', format='docx')

    latex_code = latex_code.removeprefix('\(').removesuffix('\)\n')
    print(f"latex_code: {repr(latex_code)}")

    return latex_code


@app.template_filter('urlencode')
def urlencode_filter(s):
    if type(s) == 'Markup':
        s = s.unescape()
    s = s.encode('utf8')
    s = quote(s)
    return Markup(s)


def determine_mode(omath_element, nsmap):
    # 선형 모드 패턴
    linear_mode_tags = omath_element.findall('.//m:t', nsmap)
    
    # 전문가 모드 패턴
    professional_mode_tags = omath_element.findall('.//m:f', nsmap)
    
    # 선형 모드 태그만 존재하는 경우
    if linear_mode_tags and not professional_mode_tags:
        return 'linear'
    # 전문가 모드 태그가 존재하는 경우
    elif professional_mode_tags:
        return 'professional'
    # 아무 패턴도 맞지 않는 경우 (이는 보통 발생하지 않을 것입니다)
    else:
        return 'unknown'


if __name__ == '__main__':
    app.run(debug=True)
